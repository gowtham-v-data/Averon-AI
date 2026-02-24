from fastapi import FastAPI, UploadFile, File
from pydantic import BaseModel
from fastapi.middleware.cors import CORSMiddleware
from sqlalchemy import create_engine, text
from dotenv import load_dotenv
from pinecone import Pinecone
from sentence_transformers import SentenceTransformer
from pypdf import PdfReader
import requests
import os
import uuid
import io
from gtts import gTTS
from fastapi.responses import StreamingResponse, JSONResponse
from starlette.concurrency import run_in_threadpool
import json
import re

def clean_json_response(text: str) -> str:
    """Clean LLM response to extract valid JSON, stripping markdown fences and extra text."""
    text = text.strip()
    
    # Try to extract content between ```json and ``` or ``` and ```
    # Support both case-insensitive json and optional whitespace
    match = re.search(r'```(?:json|JSON)?\s*(.*?)\s*```', text, re.DOTALL)
    if match:
        result = match.group(1).strip()
    else:
        # If no fences, try to find the outermost braces
        first_brace = text.find('{')
        last_brace = text.rfind('}')
        if first_brace != -1 and last_brace != -1 and last_brace > first_brace:
            result = text[first_brace:last_brace + 1]
        else:
            result = text
    
    return result.strip()

# Load environment variables
load_dotenv()

PINECONE_API_KEY = os.getenv("PINECONE_API_KEY")
INDEX_NAME = os.getenv("PINECONE_INDEX_NAME")
GROQ_API_KEY = os.getenv("GROQ_API_KEY")

# Validate required environment variables
if not GROQ_API_KEY:
    print("âš ï¸ WARNING: GROQ_API_KEY not found in environment variables!")
    print("   Please add it to your .env file")

if not PINECONE_API_KEY:
    print("âš ï¸ WARNING: PINECONE_API_KEY not found in environment variables!")
    
if not INDEX_NAME:
    print("âš ï¸ WARNING: PINECONE_INDEX_NAME not found in environment variables!")

print("âœ… Environment variables loaded")
print(f"   GROQ_API_KEY: {'âœ“ Set' if GROQ_API_KEY else 'âœ— Missing'}")
print(f"   PINECONE_API_KEY: {'âœ“ Set' if PINECONE_API_KEY else 'âœ— Missing'}")
print(f"   INDEX_NAME: {INDEX_NAME if INDEX_NAME else 'âœ— Missing'}")

# Initialize app
app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Initialize database
# Database connection can take time, we wrap it in a try-except
DATABASE_URL = "postgresql://postgres:gow%402005@localhost/claude_clone"
try:
    engine = create_engine(DATABASE_URL)
except Exception as e:
    print(f"Warning: Database connection failed: {e}")

# Initialize embedding model (Small and fast)
embedding_model = SentenceTransformer("all-MiniLM-L6-v2")

# Initialize Pinecone
try:
    pc = Pinecone(api_key=PINECONE_API_KEY)
    index = pc.Index(INDEX_NAME)
except Exception as e:
    print(f"Warning: Pinecone initialization failed: {e}")


class ChatRequest(BaseModel):
    message: str

class TTSRequest(BaseModel):
    text: str

# Keywords that indicate the user is asking about uploaded documents
DOC_KEYWORDS = ["pdf", "document", "file", "upload", "summarize", "summary", "resume", "attached", "knowledge base"]

# Keywords for website builder detection (includes common typos)
BUILD_KEYWORDS = ["build", "bulid", "biuld", "buid", "create", "make", "generate", "design", "develop"]
SITE_KEYWORDS = ["website", "webpage", "web page", "web site", "site", "landing page", "web app", "webapp", "homepage", "home page"]

def is_builder_request(msg: str) -> bool:
    """Check if the user wants to build a website (flexible matching with typo tolerance)."""
    msg_lower = msg.lower()
    has_build = any(kw in msg_lower for kw in BUILD_KEYWORDS)
    has_site = any(kw in msg_lower for kw in SITE_KEYWORDS)
    return has_build and has_site

# â”€â”€â”€ WEBSITE BUILDER: SYSTEM MESSAGE â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
BUILDER_SYSTEM_MSG = (
    "You are an elite frontend architect at a world-class design agency "
    "(think Vercel, Linear, Stripe quality). "
    "You generate complete, production-ready website code that could be deployed immediately. "
    "You respond ONLY with valid JSON containing a 'files' array. "
    "Each file object has 'name' (string) and 'content' (string) keys. "
    "You NEVER include markdown formatting, explanations, backticks, or commentary â€” ONLY the raw JSON object. "
    "Your websites are visually stunning with premium aesthetics, fully responsive, accessible, "
    "and feature smooth modern animations. "
    "Every page you build looks like it was crafted by a senior developer with 10+ years of experience."
)

# â”€â”€â”€ WEBSITE BUILDER: PROMPT GENERATOR â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€â”€
def get_builder_prompt(user_message: str) -> str:
    """Generate the professional builder prompt for the LLM."""
    return (
        f"PROJECT BRIEF: {user_message}\n"
        "\n"
        "INSTRUCTION: Return ONLY a valid JSON object. No markdown, no commentary.\n"
        "FORMAT: {\"files\": [{\"name\": \"index.html\", \"content\": \"...\"}, {\"name\": \"style.css\", \"content\": \"...\"}]}\n"
        "\n"
        "ARCHITECTURE:\n"
        "\n"
        "â”â”â”â”â”â”â” HTML ARCHITECTURE â”â”â”â”â”â”â”\n"
        'â€¢ DOCTYPE html with lang="en"\n'
        "â€¢ <head>: charset UTF-8, viewport meta, unique <title>, meta description, Open Graph tags, link to style.css\n"
        'â€¢ Google Fonts: Import "Inter" font (weights 300,400,500,600,700) via <link> tag â€” use as primary font-family\n'
        "â€¢ Semantic HTML5: <header>, <nav>, <main>, <section>, <article>, <footer>\n"
        "â€¢ Every page shares the SAME <header> with logo + responsive navigation + mobile hamburger icon\n"
        "â€¢ Every page shares the SAME <footer> with organized link columns, social icons (Unicode/SVG), copyright with auto-year\n"
        "â€¢ Navigation highlights the current active page with a CSS class\n"
        'â€¢ All internal links work correctly (href="about.html" etc.)\n'
        "â€¢ Add smooth scroll: html { scroll-behavior: smooth }\n"
        "â€¢ script.js loaded with defer attribute\n"
        "\n"
        "â”â”â”â”â”â”â” PAGE DETAILS â”â”â”â”â”â”â”\n"
        "\n"
        "INDEX.HTML â€” Home Page:\n"
        "â€¢ Hero section: large bold heading, compelling subtitle, gradient or overlay background, 1-2 CTA buttons with hover effects\n"
        "â€¢ Features/services grid: 3-4 cards with Unicode icons, title, description\n"
        "â€¢ Social proof section: testimonials carousel OR animated stats/counters\n"
        "â€¢ Call-to-action banner with gradient background before footer\n"
        "\n"
        "ABOUT.HTML â€” About Page:\n"
        "â€¢ Hero banner with page title and breadcrumb\n"
        "â€¢ Brand story section with side-by-side layout (text + visual placeholder)\n"
        "â€¢ Mission / Vision / Values in a 3-column grid with icons\n"
        "â€¢ Team section OR timeline/milestones section\n"
        "\n"
        "SERVICES.HTML â€” Services/Products Page:\n"
        "â€¢ Hero banner with page title\n"
        "â€¢ Service/product cards in responsive grid (icon, title, description, price/CTA)\n"
        "â€¢ FAQ accordion section with working JavaScript expand/collapse and smooth height animation\n"
        "â€¢ Comparison table or pricing tiers if relevant\n"
        "\n"
        "CONTACT.HTML â€” Contact Page:\n"
        "â€¢ Hero banner with page title\n"
        "â€¢ Contact form: name, email, phone, subject <select>, message <textarea>, submit button\n"
        "â€¢ Full form validation with JavaScript + visual error/success states\n"
        "â€¢ Contact info cards with icons\n"
        "â€¢ Business hours or additional info section\n"
        "\n"
        "â”â”â”â”â”â”â” CSS DESIGN SYSTEM â”â”â”â”â”â”â”\n"
        "â€¢ CSS Custom Properties (--variables) for: colors (primary, secondary, accent, bg, text, border), spacing scale, border-radius, shadows, font-sizes\n"
        "â€¢ Professional color palette: dark rich backgrounds (#0a0a0a, #111827, #1e293b) with vibrant accents â€” OR a light clean theme with bold accent colors. Follow the 60-30-10 color rule.\n"
        "â€¢ Typography: font-family 'Inter', sans-serif. Use clamp() for fluid responsive font sizes. Proper hierarchy (h1 > h2 > h3 > body).\n"
        "â€¢ Layered box-shadows for depth: --shadow-sm, --shadow-md, --shadow-lg, --shadow-xl\n"
        "â€¢ Gradient backgrounds on hero sections using linear-gradient or radial-gradient with beautiful color stops\n"
        "â€¢ Glassmorphism where appropriate: backdrop-filter: blur(12px); background: rgba(255,255,255,0.05); border: 1px solid rgba(255,255,255,0.1)\n"
        "â€¢ Smooth transitions on ALL interactive elements: transition: all 0.3s cubic-bezier(0.4, 0, 0.2, 1)\n"
        "â€¢ Button styles: primary (filled gradient), secondary (outlined), ghost â€” all with hover scale + shadow elevation + color shift\n"
        "â€¢ Card hover effect: transform: translateY(-4px) with enhanced shadow\n"
        "â€¢ Sticky header with backdrop-filter blur, gains shadow on scroll via JS class .scrolled\n"
        "â€¢ Mobile hamburger menu: smooth slide-in or dropdown animation\n"
        "â€¢ Responsive: mobile-first, breakpoints at 768px and 1024px using @media\n"
        "â€¢ @keyframes for: fadeInUp, fadeInLeft, fadeInRight, scaleIn â€” used with .visible class\n"
        "â€¢ Form inputs: custom styled with focus ring (box-shadow outline), error state (red border), success state (green border)\n"
        "â€¢ Footer: darker background, multi-column grid layout\n"
        "â€¢ Smooth scrollbar styling (webkit-scrollbar)\n"
        "â€¢ Selection color styling (::selection)\n"
        "\n"
        "â”â”â”â”â”â”â” JAVASCRIPT FEATURES â”â”â”â”â”â”â”\n"
        "â€¢ DOMContentLoaded wrapper for all code\n"
        "â€¢ Mobile menu: hamburger toggle with classList.toggle, animate open/close\n"
        "â€¢ Sticky header: window scroll listener â€” add .scrolled class when scrollY > 50\n"
        'â€¢ Smooth scroll for all anchor links (querySelectorAll(\'a[href^="#"]\'))\n'
        "â€¢ Active nav highlighting: compare window.location with href, add .active class\n"
        "â€¢ Scroll-reveal animations: IntersectionObserver on elements with [data-animate] â€” add .visible class\n"
        "â€¢ FAQ accordion: querySelectorAll('.faq-item'), click toggles .active, smooth max-height transition\n"
        "â€¢ Contact form: addEventListener('submit'), validate all fields, display inline error messages, show success message on valid submit (preventDefault)\n"
        "â€¢ Stats counter: IntersectionObserver triggers countUp animation (setInterval from 0 to target)\n"
        "â€¢ Back-to-top button: appears when scrollY > 300, smooth scroll to top on click\n"
        "â€¢ Auto-update copyright year: textContent = new Date().getFullYear()\n"
        "â€¢ NO external libraries â€” 100% vanilla JavaScript\n"
        "â€¢ Clean modular code with descriptive function names and comments\n"
        "\n"
        "â”â”â”â”â”â”â” QUALITY STANDARDS â”â”â”â”â”â”â”\n"
        "â€¢ ZERO placeholder text â€” all content realistic and contextual to the project brief\n"
        "â€¢ ZERO broken links â€” every href points to a valid page\n"
        "â€¢ Accessibility: alt on images, ARIA labels on buttons/nav, :focus-visible outlines, prefers-reduced-motion media query\n"
        "â€¢ SEO: unique <title> and <meta description> per page, single <h1> per page, proper heading hierarchy\n"
        "â€¢ Clean indented code with meaningful BEM-inspired class names (e.g., .hero__title, .card--featured)\n"
        "â€¢ Production-ready â€” deployable as-is to any static hosting\n"
        "â€¢ Output ONLY the JSON object â€” no markdown, no explanations, no backticks"
    )


@app.post("/chat")
def chat(req: ChatRequest):
    user_message = req.message.strip()
    
    # Check if API key is configured
    if not GROQ_API_KEY:
        return {"reply": "âš ï¸ Server configuration error: GROQ API key is missing. Please contact the administrator."}

    # WEBSITE BUILDER MODE
    if is_builder_request(user_message):
        builder_prompt = get_builder_prompt(user_message)
        print(f"ðŸ”¨ Website builder triggered for: {user_message[:80]}...")

        try:
            response = requests.post(
                "https://api.groq.com/openai/v1/chat/completions",
                headers={
                    "Authorization": f"Bearer {GROQ_API_KEY}",
                    "Content-Type": "application/json"
                },
                json={
                    "model": "llama-3.3-70b-versatile",
                    "messages": [
                        {"role": "system", "content": BUILDER_SYSTEM_MSG},
                        {"role": "user", "content": builder_prompt}
                    ],
                    "temperature": 0.4,
                    "max_tokens": 8192,
                    "response_format": {"type": "json_object"}
                },
                timeout=120
            )
            result = response.json()
            
            print(f"Website builder API response status: {response.status_code}")
            
            if "error" in result:
                error_msg = result["error"].get("message", "Unknown error")
                print(f"Website builder API error: {error_msg}")
                return {"reply": f"Error generating website: {error_msg}"}
            
            if "choices" not in result:
                print(f"Unexpected website builder response: {result}")
                return {"reply": "Error generating website. Please try again."}
            
            raw_content = result["choices"][0]["message"]["content"]
            print(f"Website builder raw response length: {len(raw_content)} chars")
            
            # Clean up the response: strip markdown code fences and extra text
            cleaned = clean_json_response(raw_content)
            
            # Validate it's actually parseable JSON
            try:
                parsed = json.loads(cleaned)
                
                # Handle case where files might be nested differently
                if "files" not in parsed:
                    # Try to find files key in any nested structure
                    print(f"JSON keys received: {list(parsed.keys())}")
                    return {"reply": "Error: AI generated invalid website structure. Please try again with a simpler request."}
                
                # Validate each file has name and content
                for f in parsed["files"]:
                    if "name" not in f or "content" not in f:
                        print(f"Invalid file structure: {list(f.keys())}")
                        return {"reply": "Error: AI generated incomplete files. Please try again."}
                
                print(f"âœ… Website built successfully! {len(parsed['files'])} files generated.")
                for f in parsed["files"]:
                    print(f"   ðŸ“„ {f['name']} ({len(f['content'])} chars)")
                
                # Return the cleaned, validated JSON string
                return {"builder": json.dumps(parsed)}
                
            except json.JSONDecodeError as je:
                print(f"JSON parse error after cleanup: {je}")
                print(f"Raw content (first 500 chars): {raw_content[:500]}")
                return {"reply": "Error: AI response was not valid JSON. Please try again."}
            
        except requests.exceptions.Timeout:
            print("Website builder request timed out")
            return {"reply": "Website generation timed out. Please try a simpler request."}
        except Exception as e:
            print(f"Website builder error: {e}")
            return {"reply": f"Error generating website: {str(e)}"}

    # NORMAL RAG CHAT MODE
    # 1. History
    try:
        with engine.connect() as conn:
            res = conn.execute(text("SELECT role, content FROM messages ORDER BY created_at DESC LIMIT 10"))
            history = [{"role": row[0], "content": row[1]} for row in res]
            history.reverse()
    except Exception as e:
        print(f"âš ï¸ History fetch error: {e}")
        history = []

    # 2. Check if user is asking about uploaded documents
    is_doc_query = any(kw in user_message.lower() for kw in DOC_KEYWORDS)
    
    # 3. Get list of uploaded files for context
    uploaded_filenames = []
    try:
        with engine.connect() as conn:
            res = conn.execute(text("SELECT filename FROM uploaded_files ORDER BY uploaded_at DESC"))
            uploaded_filenames = [row[0] for row in res]
    except Exception as e:
        print(f"âš ï¸ File list fetch error: {e}")

    # 4. Vector Search - use more results for document queries
    context = ""
    try:
        query_embedding = embedding_model.encode(user_message).tolist()
        # Fetch more chunks when user asks about documents
        top_k = 15 if is_doc_query else 5
        results = index.query(vector=query_embedding, top_k=top_k, include_metadata=True)
        
        if results["matches"]:
            context = "\n".join([m["metadata"]["text"] for m in results["matches"]])
            print(f"âœ… Vector search returned {len(results['matches'])} results (top_k={top_k})")
            print(f"   Context preview: {context[:200]}...")
        else:
            print("âš ï¸ Vector search returned 0 matches")
            context = ""
    except Exception as e:
        print(f"âŒ Vector search error: {e}")
        context = ""

    # 5. Build system prompt with knowledge base awareness
    system_prompt = "You are Averon, a powerful and helpful AI assistant. You are friendly, knowledgeable, and provide clear, well-structured responses."
    if uploaded_filenames:
        file_list = ", ".join(uploaded_filenames)
        system_prompt += f"\n\nThe user has uploaded the following documents to the knowledge base: {file_list}."
        system_prompt += "\nYou have access to the content of these documents through the context provided below."
        system_prompt += "\nWhen the user asks about their uploaded documents, PDFs, or files, use the provided context to answer."
        system_prompt += "\nIf the user asks to summarize a document, provide a thorough summary based on the context."

    # 6. LLM Call
    messages = [
        {"role": "system", "content": system_prompt},
    ]
    
    # Only add context if we actually have some
    if context.strip():
        messages.append({"role": "system", "content": f"Content from the user's uploaded documents:\n{context}"})
    
    messages += history
    messages.append({"role": "user", "content": user_message})

    try:
        response = requests.post(
            "https://api.groq.com/openai/v1/chat/completions",
            headers={"Authorization": f"Bearer {GROQ_API_KEY}"},
            json={"model": "llama-3.1-8b-instant", "messages": messages},
            timeout=30
        )
        result = response.json()
        
        # Check if API returned an error
        if "error" in result:
            error_msg = result["error"].get("message", "Unknown API error")
            print(f"API Error: {error_msg}")
            return {"reply": f"Sorry, I encountered an error: {error_msg}"}
        
        # Check if response has expected format
        if "choices" not in result:
            print(f"Unexpected API response: {result}")
            return {"reply": "Sorry, I received an unexpected response from the AI service. Please try again."}
        
        reply = result["choices"][0]["message"]["content"]
        
    except requests.exceptions.Timeout:
        print("API request timed out")
        return {"reply": "Sorry, the request timed out. Please try again."}
    except requests.exceptions.RequestException as e:
        print(f"API request failed: {e}")
        return {"reply": "Sorry, I'm having trouble connecting to the AI service. Please try again later."}
    except Exception as e:
        print(f"Unexpected error: {e}")
        return {"reply": "Sorry, an unexpected error occurred. Please try again."}

    # 7. Save
    try:
        with engine.connect() as conn:
            conn.execute(text("INSERT INTO messages (role, content) VALUES (:r, :c)"), {"r": "user", "c": user_message})
            conn.execute(text("INSERT INTO messages (role, content) VALUES (:r, :c)"), {"r": "assistant", "c": reply})
            conn.commit()
    except Exception as e:
        print(f"âš ï¸ Message save error: {e}")

    return {"reply": reply}

@app.post("/tts")
async def text_to_speech(req: TTSRequest):
    """Generate speech audio from text using Google TTS"""
    print(f"🔊 TTS Request: {req.text[:50]}...")
    if not req.text.strip():
        return {"error": "No text provided"}
    
    try:
        # Run gTTS in a threadpool to avoid blocking the event loop
        def generate_audio():
            tts = gTTS(text=req.text, lang='en', slow=False)
            audio_buffer = io.BytesIO()
            tts.write_to_fp(audio_buffer)
            audio_buffer.seek(0)
            return audio_buffer

        audio_buffer = await run_in_threadpool(generate_audio)
        
        return StreamingResponse(
            audio_buffer,
            media_type="audio/mpeg",
            headers={
                "Content-Disposition": "inline; filename=speech.mp3",
                "Cache-Control": "no-cache"
            }
        )
    except Exception as e:
        print(f"❌ TTS Error: {e}")
        return JSONResponse(status_code=500, content={"error": str(e)})

@app.post("/upload")
async def upload_document(file: UploadFile = File(...)):
    filename = file.filename.lower()
    original_filename = file.filename  # Keep original name for display
    file_text = ""

    print(f"ðŸ“„ Upload request received: {filename}")

    try:
        file_bytes = await file.read()

        if filename.endswith(".pdf"):
            reader = PdfReader(io.BytesIO(file_bytes))
            file_text = "\n".join([p.extract_text() for p in reader.pages if p.extract_text()])
            print(f"   PDF pages: {len(reader.pages)}, extracted text length: {len(file_text)} chars")

        elif filename.endswith(".csv"):
            import pandas as pd
            df = pd.read_csv(io.BytesIO(file_bytes))
            file_text = df.to_string()
            print(f"   CSV rows: {len(df)}, columns: {list(df.columns)}")

        elif filename.endswith((".xlsx", ".xls")):
            import pandas as pd
            df = pd.read_excel(io.BytesIO(file_bytes))
            file_text = df.to_string()
            print(f"   Excel rows: {len(df)}, columns: {list(df.columns)}")

        elif filename.endswith(".docx"):
            from docx import Document
            doc = Document(io.BytesIO(file_bytes))
            file_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
            print(f"   DOCX paragraphs: {len(doc.paragraphs)}")

        elif filename.endswith(".doc"):
            # .doc files need different handling - try to read as text
            try:
                file_text = file_bytes.decode('utf-8', errors='ignore')
            except:
                file_text = file_bytes.decode('latin-1', errors='ignore')
            print(f"   DOC text length: {len(file_text)} chars")

        elif filename.endswith(".txt"):
            file_text = file_bytes.decode('utf-8', errors='ignore')
            print(f"   TXT text length: {len(file_text)} chars")

        elif filename.endswith(".json"):
            file_text = file_bytes.decode('utf-8', errors='ignore')
            print(f"   JSON text length: {len(file_text)} chars")

        else:
            # Try to read as plain text for any other file
            try:
                file_text = file_bytes.decode('utf-8', errors='ignore')
                print(f"   Generic file read as text: {len(file_text)} chars")
            except:
                return {"error": f"Unsupported file type: {filename.split('.')[-1]}"}

    except Exception as e:
        print(f"   âŒ File read error: {e}")
        return {"error": f"Error reading file: {str(e)}"}

    if not file_text.strip():
        print(f"   âš ï¸ No text extracted from {filename}")
        return {"error": "Empty file or no text could be extracted."}

    print(f"   ðŸ“ Text preview: {file_text[:200]}...")

    # Vectorize
    chunks = [file_text[i:i+500] for i in range(0, len(file_text), 500)]
    vectors = []
    for chunk in chunks:
        vectors.append({
            "id": str(uuid.uuid4()),
            "values": embedding_model.encode(chunk).tolist(),
            "metadata": {"text": chunk, "filename": filename}
        })
    index.upsert(vectors)
    print(f"   âœ… Upserted {len(vectors)} chunks to Pinecone")

    # Database Log
    try:
        with engine.connect() as conn:
            conn.execute(text("INSERT INTO uploaded_files (filename) VALUES (:f)"), {"f": original_filename})
            conn.commit()
        print(f"   âœ… Logged {original_filename} to database")
    except Exception as e:
        print(f"   âš ï¸ DB log error: {e}")

    return {"status": f"{original_filename} uploaded successfully!"}

@app.get("/files")
def get_files():
    try:
        with engine.connect() as conn:
            res = conn.execute(text("SELECT id, filename, uploaded_at FROM uploaded_files ORDER BY uploaded_at DESC"))
            return {"files": [{"id": r[0], "filename": r[1], "uploaded_at": str(r[2])} for r in res]}
    except:
        return {"files": []}

@app.post("/clear-chat")
def clear_chat():
    """Clear all chat messages from the database for a fresh conversation."""
    try:
        with engine.connect() as conn:
            conn.execute(text("DELETE FROM messages"))
            conn.commit()
        print("ðŸ§¹ Chat history cleared")
        return {"status": "Chat history cleared."}
    except Exception as e:
        print(f"âš ï¸ Clear chat error: {e}")
        return {"error": str(e)}

@app.post("/clear-docs")
def clear_documents():
    try:
        if index:
            index.delete(delete_all=True)
        with engine.connect() as conn:
            conn.execute(text("DELETE FROM uploaded_files"))
            conn.commit()
        return {"status": "Cleared."}
    except Exception as e:
        print(f"Clear docs error: {e}")
        return {"error": str(e)}

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="127.0.0.1", port=8000)
